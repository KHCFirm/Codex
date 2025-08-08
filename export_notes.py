from __future__ import annotations

import csv
from collections import defaultdict
from pathlib import Path
from typing import Dict, List

from docx import Document
from docx.shared import Pt
from fpdf import FPDF

ALLOWED_TYPES = {
    "Personal Injury",
    "Workers' Compensation",
    "Master Multi - PI & WC",
}


def group_rows(csv_path: Path) -> Dict[str, List[dict]]:
    groups: Dict[str, List[dict]] = defaultdict(list)
    with open(csv_path, newline="", encoding="utf-8") as fh:
        reader = csv.DictReader(fh)
        for row in reader:
            if row.get("Project Type") not in ALLOWED_TYPES:
                continue
            pid = row.get("Project ID")
            groups[pid].append(row)
    return groups


def write_docx(rows: List[dict], path: Path) -> None:
    doc = Document()
    for row in rows:
        lines = [
            f"Project ID: {row['Project ID']}",
            f"Project Name: {row['Project Name']}",
            f"Project Type: {row['Project Type']}",
            f"{row['Created At']} - {row['Author Name']}: {row['Note Text']}",
            "",
        ]
        for line in lines:
            p = doc.add_paragraph()
            run = p.add_run(line)
            run.font.size = Pt(8)
    doc.save(path)


def write_pdf(rows: List[dict], path: Path) -> None:
    pdf = FPDF()
    pdf.set_auto_page_break(auto=True, margin=15)
    pdf.add_page()
    pdf.set_font("Helvetica", size=8)
    for row in rows:
        pdf.multi_cell(0, 4, f"Project ID: {row['Project ID']}")
        pdf.multi_cell(0, 4, f"Project Name: {row['Project Name']}")
        pdf.multi_cell(0, 4, f"Project Type: {row['Project Type']}")
        pdf.ln(1)
        pdf.multi_cell(
            0,
            4,
            f"{row['Created At']} - {row['Author Name']}: {row['Note Text']}",
        )
        pdf.ln(4)
    pdf.output(str(path))


def export(csv_path: str | Path, out_dir: str | Path) -> None:
    csv_path = Path(csv_path)
    out_dir = Path(out_dir)
    out_dir.mkdir(parents=True, exist_ok=True)
    groups = group_rows(csv_path)
    for pid, rows in groups.items():
        write_docx(rows, out_dir / f"{pid}.docx")
        write_pdf(rows, out_dir / f"{pid}.pdf")


if __name__ == "__main__":
    import argparse

    parser = argparse.ArgumentParser(
        description="Convert project notes spreadsheet into Word and PDF files."
    )
    parser.add_argument("csv_path", help="Path to the input CSV file")
    parser.add_argument(
        "out_dir", help="Directory where the Word and PDF files will be written"
    )
    args = parser.parse_args()
    export(args.csv_path, args.out_dir)
